import streamlit as st
import fitz  # PyMuPDF
import chardet
import io
import zipfile
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import os
from PIL import Image, ImageFont, ImageDraw, ImageEnhance
import pytesseract
import numpy as np
import cv2
import google.generativeai as genai
import base64
from typing import List, Dict, Optional, Tuple
import time
import logging
import asyncio
import concurrent.futures
from threading import Lock
import gc
import psutil
import threading
from queue import Queue
import json
from datetime import datetime

# Set page configuration
st.set_page_config(
    page_title="Enhanced Large PDF Text Extractor with AI",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Enhanced OCR configuration for Indic languages
INDIC_LANGUAGES = {
    'hindi': 'hin',
    'sanskrit': 'san',
    'bengali': 'ben',
    'gujarati': 'guj',
    'kannada': 'kan',
    'malayalam': 'mal',
    'marathi': 'mar',
    'punjabi': 'pan',
    'tamil': 'tam',
    'telugu': 'tel',
    'urdu': 'urd',
    'english': 'eng'
}

# Enhanced Indic font families that work well with OCR
INDIC_FONTS = {
    'devanagari': ['Mangal', 'Nirmala UI', 'Sanskrit 2003', 'Kokila', 'Aparajita'],
    'bengali': ['Nirmala UI', 'Shonar Bangla', 'Vrinda'],
    'gujarati': ['Nirmala UI', 'Shruti'],
    'kannada': ['Nirmala UI', 'Tunga'],
    'malayalam': ['Nirmala UI', 'Kartika'],
    'tamil': ['Nirmala UI', 'Latha'],
    'telugu': ['Nirmala UI', 'Gautami'],
    'punjabi': ['Nirmala UI', 'Raavi']
}

# Global variables for rate limiting and memory management
gemini_lock = Lock()
last_gemini_call = 0
processed_pages_cache = {}
memory_threshold = 80  # Percentage


class EnhancedRateLimiter:
    """Enhanced rate limiter for Gemini API with accuracy-focused adaptive delays"""

    def __init__(self, calls_per_minute=20, burst_limit=3, accuracy_mode=False):
        # Slower for accuracy
        self.calls_per_minute = calls_per_minute if not accuracy_mode else 15
        self.burst_limit = burst_limit if not accuracy_mode else 2
        self.call_times = []
        self.lock = Lock()
        self.consecutive_calls = 0
        self.last_call_time = 0
        self.accuracy_mode = accuracy_mode

    def wait_if_needed(self):
        with self.lock:
            now = time.time()

            # Remove calls older than 1 minute
            self.call_times = [t for t in self.call_times if now - t < 60]

            # Enhanced adaptive delay for accuracy
            base_delay = 2.0 if self.accuracy_mode else 1.0
            if self.consecutive_calls > 0:
                adaptive_delay = min(
                    5.0, base_delay * (1 + self.consecutive_calls / 20))
                if now - self.last_call_time < adaptive_delay:
                    time.sleep(adaptive_delay - (now - self.last_call_time))

            # Check if we need to wait for rate limit
            if len(self.call_times) >= self.calls_per_minute:
                sleep_time = 60 - (now - self.call_times[0])
                if sleep_time > 0:
                    time.sleep(sleep_time)

            # Check burst limit with stricter enforcement for accuracy
            recent_window = 15 if self.accuracy_mode else 10
            recent_calls = [
                t for t in self.call_times if now - t < recent_window]
            if len(recent_calls) >= self.burst_limit:
                time.sleep(3 if self.accuracy_mode else 2)

            self.call_times.append(now)
            self.consecutive_calls += 1
            self.last_call_time = now

    def reset_consecutive(self):
        self.consecutive_calls = 0


# Global rate limiter
rate_limiter = EnhancedRateLimiter(accuracy_mode=True)


def get_memory_usage():
    """Get current memory usage percentage"""
    try:
        return psutil.virtual_memory().percent
    except:
        return 0


def cleanup_memory():
    """Force garbage collection and cleanup"""
    gc.collect()
    if hasattr(gc, 'set_threshold'):
        gc.set_threshold(700, 10, 10)


def setup_gemini(api_key: str) -> bool:
    """Setup Gemini API with user's API key"""
    try:
        genai.configure(api_key=api_key)
        # Test the API key with a simple request
        model = genai.GenerativeModel('gemini-2.0-flash-exp')
        test_response = model.generate_content("Test connection",
                                               generation_config=genai.types.GenerationConfig(
                                                   temperature=0.1,
                                                   max_output_tokens=10
                                               ))
        return True
    except Exception as e:
        st.error(f"Failed to setup Gemini API: {str(e)}")
        return False


def enhanced_image_optimization_for_gemini(image: Image.Image, quality_mode='high'):
    """Enhanced image optimization specifically for maximum accuracy"""
    try:
        # Convert to RGB if necessary
        if image.mode in ('RGBA', 'LA', 'P'):
            rgb_image = Image.new('RGB', image.size, (255, 255, 255))
            if image.mode == 'P':
                image = image.convert('RGB')
            else:
                rgb_image.paste(image, mask=image.split()
                                [-1] if image.mode in ('RGBA', 'LA') else None)
                image = rgb_image

        # High-quality mode: maintain higher resolution
        if quality_mode == 'high':
            max_size = (3072, 3072)  # Increased from 2048
            quality = 95  # Increased from 85
        else:
            max_size = (2048, 2048)
            quality = 85

        # Resize more conservatively
        if image.size[0] > max_size[0] or image.size[1] > max_size[1]:
            # Use high-quality resampling
            image.thumbnail(max_size, Image.Resampling.LANCZOS)

        # Apply image enhancement for better text recognition
        # Contrast enhancement
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.2)

        # Sharpness enhancement
        enhancer = ImageEnhance.Sharpness(image)
        image = enhancer.enhance(1.1)

        # Convert to bytes with high quality
        img_byte_arr = io.BytesIO()
        image.save(img_byte_arr, format='JPEG', optimize=True, quality=quality)
        img_byte_arr.seek(0)

        return Image.open(img_byte_arr)

    except Exception as e:
        st.warning(f"Image optimization failed: {str(e)}, using original")
        return image


def extract_text_with_gemini_accuracy_focused(images_data: List[Tuple[int, Image.Image]],
                                              selected_languages: List[str],
                                              extraction_type: str = "general") -> Dict[int, Tuple[str, float]]:
    """Extract text using Gemini with maximum accuracy focus"""
    results = {}

    try:
        model = genai.GenerativeModel('gemini-2.0-flash-exp')

        # Language setup with better descriptions
        language_names = {
            'hindi': 'Hindi (हिंदी) - Devanagari script',
            'sanskrit': 'Sanskrit (संस्कृत) - Devanagari script',
            'bengali': 'Bengali (বাংলা) - Bengali script',
            'gujarati': 'Gujarati (ગુજરાતી) - Gujarati script',
            'kannada': 'Kannada (ಕನ್ನಡ) - Kannada script',
            'malayalam': 'Malayalam (മലയാളം) - Malayalam script',
            'marathi': 'Marathi (मराठी) - Devanagari script',
            'punjabi': 'Punjabi (ਪੰਜਾਬੀ) - Gurmukhi script',
            'tamil': 'Tamil (தமிழ்) - Tamil script',
            'telugu': 'Telugu (తెలుగు) - Telugu script',
            'urdu': 'Urdu (اردو) - Arabic script',
            'english': 'English - Latin script'
        }

        selected_lang_names = [language_names.get(
            lang, lang) for lang in selected_languages]

        # Enhanced prompts for maximum accuracy
        if extraction_type == "academic":
            base_prompt = f"""You are an expert document digitization specialist. Extract ALL text from this academic document with absolute precision:

Target Languages: {', '.join(selected_lang_names)}

CRITICAL REQUIREMENTS:
1. Extract EVERY visible character, symbol, number, and punctuation mark
2. Preserve exact formatting: line breaks, spacing, indentation, bullet points
3. Maintain mathematical formulas, equations, and special symbols exactly as shown
4. Preserve table structures, headers, and footnotes
5. Include ALL references, citations, and bibliographic information
6. Capture figure captions, table titles, and margin notes
7. Maintain original text hierarchy (headings, subheadings, paragraphs)

ACCURACY STANDARDS:
- Zero tolerance for missing text
- Perfect character-level accuracy required
- Preserve all diacritical marks and special characters
- Maintain exact punctuation and spacing

Output the complete, unmodified text exactly as it appears in the document."""

        elif extraction_type == "indic_specialized":
            base_prompt = f"""You are a master specialist in Indian language document processing. Extract text with cultural and linguistic expertise:

Target Languages: {', '.join(selected_lang_names)}

EXPERT REQUIREMENTS:
1. Perfect recognition of complex Indic scripts and conjunct characters
2. Accurate handling of diacritical marks, matras, and combining characters  
3. Proper word segmentation respecting linguistic boundaries
4. Cultural context awareness for proper noun recognition
5. Accurate transliteration preservation where present
6. Perfect handling of Sanskrit verses, religious texts, and cultural terminology
7. Precise recognition of regional script variations

SCRIPT-SPECIFIC EXPERTISE:
- Devanagari: Perfect conjunct recognition, proper matra placement
- Bengali: Accurate handling of complex ligatures and vowel marks
- Tamil: Proper recognition of agglutinative word structures
- Telugu/Kannada: Correct handling of vowel-consonant combinations
- Arabic (Urdu): Right-to-left text flow, proper diacritic marks

Extract with 100% cultural and linguistic accuracy."""

        elif extraction_type == "handwritten":
            base_prompt = f"""You are an expert paleographer specializing in handwritten document analysis:

Target Languages: {', '.join(selected_lang_names)}

HANDWRITING ANALYSIS REQUIREMENTS:
1. Carefully analyze individual character formations and writing style
2. Consider contextual clues for ambiguous characters
3. Handle cursive connections and character variations
4. Recognize personal writing idiosyncrasies and abbreviations
5. Use linguistic context to resolve unclear characters
6. Maintain original punctuation and formatting choices
7. Preserve author's spacing and paragraph structures

ACCURACY APPROACH:
- Character-by-character careful analysis
- Cross-reference unclear portions with context
- Mark genuinely unclear sections with [UNCLEAR] rather than guessing
- Maintain original capitalization and punctuation patterns
- Preserve personal formatting choices and emphasis

Extract with scholarly precision, noting any uncertainties clearly."""

        else:  # general - enhanced for accuracy
            base_prompt = f"""You are a professional document digitization expert. Extract ALL text with maximum accuracy:

Target Languages: {', '.join(selected_lang_names)}

PRECISION REQUIREMENTS:
1. Extract EVERY visible text element without exception
2. Maintain exact formatting: paragraphs, line breaks, spacing
3. Preserve ALL punctuation marks, symbols, and special characters
4. Include headers, footers, page numbers, and watermarks
5. Capture table contents with proper structure
6. Include figure captions, labels, and annotations
7. Maintain original text hierarchy and organization

QUALITY STANDARDS:
- Perfect character-level accuracy
- Zero missing text tolerance
- Exact punctuation preservation
- Original formatting maintenance
- Complete content capture

Provide the complete, unaltered text exactly as shown in the document."""

        # Process images individually with enhanced accuracy measures
        for page_num, image in images_data:
            try:
                # Enhanced rate limiting for accuracy
                rate_limiter.wait_if_needed()

                # High-quality image optimization
                optimized_image = enhanced_image_optimization_for_gemini(
                    image, 'high')

                # Memory check
                if get_memory_usage() > memory_threshold:
                    cleanup_memory()

                # Enhanced generation config for accuracy
                generation_config = genai.types.GenerationConfig(
                    temperature=0.05,  # Lower temperature for consistency
                    top_p=0.9,        # Slightly higher for better coverage
                    top_k=40,         # Add top_k for more focused responses
                    max_output_tokens=8192,  # Increased token limit
                    candidate_count=1
                )

                # Make API call with retry logic for accuracy
                max_retries = 3
                best_response = None
                best_length = 0

                for attempt in range(max_retries):
                    try:
                        response = model.generate_content(
                            [base_prompt, optimized_image],
                            generation_config=generation_config
                        )

                        if response.text and len(response.text.strip()) > best_length:
                            best_response = response.text
                            best_length = len(response.text.strip())

                        # If we get a good response, break early
                        if best_length > 100:
                            break

                    except Exception as retry_error:
                        if attempt == max_retries - 1:
                            raise retry_error
                        time.sleep(2 * (attempt + 1))  # Progressive backoff

                extracted_text = best_response if best_response else ""

                # Enhanced confidence calculation
                confidence = calculate_enhanced_confidence(
                    extracted_text, selected_languages, image.size)

                results[page_num] = (extracted_text, confidence)

                # Enhanced logging for accuracy tracking
                if extracted_text:
                    st.write(
                        f"✅ Page {page_num + 1}: {len(extracted_text)} characters extracted (confidence: {confidence:.1f}%)")
                else:
                    st.warning(f"⚠️ Page {page_num + 1}: No text extracted")

                # Clean up
                del optimized_image
                if get_memory_usage() > memory_threshold:
                    cleanup_memory()

            except Exception as e:
                st.error(
                    f"Gemini extraction failed for page {page_num + 1}: {str(e)}")
                results[page_num] = ("", 0)

                # Enhanced error handling for rate limits
                if "rate limit" in str(e).lower() or "quota" in str(e).lower():
                    st.warning("Rate limit hit, implementing longer delay...")
                    time.sleep(10)
                    rate_limiter.consecutive_calls += 5

    except Exception as e:
        st.error(f"Gemini batch extraction failed: {str(e)}")

    return results


def calculate_enhanced_confidence(text: str, languages: List[str], image_size: tuple) -> float:
    """Calculate confidence score with enhanced accuracy metrics"""
    if not text or not text.strip():
        return 0.0

    base_confidence = 85.0

    # Text length factor
    text_length = len(text.strip())
    if text_length > 1000:
        base_confidence += 10.0
    elif text_length > 500:
        base_confidence += 5.0
    elif text_length < 50:
        base_confidence -= 30.0

    # Language-specific bonus
    if any(lang in ['hindi', 'sanskrit', 'bengali', 'tamil'] for lang in languages):
        # Check for Indic characters
        indic_char_count = sum(1 for char in text if ord(
            char) > 2304 and ord(char) < 3071)
        if indic_char_count > 0:
            base_confidence += 8.0

    # Structure indicators (good formatting suggests accurate extraction)
    structure_score = 0
    if '\n\n' in text:  # Paragraph breaks
        structure_score += 2
    if any(char in text for char in '.,;:!?'):  # Punctuation
        structure_score += 3
    if any(char.isupper() for char in text):  # Capital letters
        structure_score += 2

    base_confidence += min(structure_score, 7)

    # Penalize very short extractions relative to image size
    pixel_count = image_size[0] * image_size[1]
    if pixel_count > 1000000 and text_length < 100:  # Large image, little text
        base_confidence -= 20.0

    return max(10.0, min(98.0, base_confidence))


def enhanced_preprocess_image_for_indic_ocr(image, language_script='devanagari'):
    """Enhanced image preprocessing with multiple techniques for maximum accuracy"""
    try:
        # Convert PIL to OpenCV format
        img_array = np.array(image)

        # Convert to grayscale if not already
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array

        # Apply multiple preprocessing techniques and choose the best
        processed_versions = []

        # Method 1: Enhanced CLAHE with script-specific parameters
        if language_script in ['devanagari', 'bengali', 'gujarati']:
            clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(gray)
            # Slight blur to connect broken characters in complex scripts
            blurred = cv2.GaussianBlur(enhanced, (1, 1), 0)
            thresh1 = cv2.adaptiveThreshold(blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                            cv2.THRESH_BINARY, 11, 2)
            processed_versions.append(('clahe_adaptive', thresh1))

        # Method 2: Otsu thresholding with enhancement
        clahe = cv2.createCLAHE(clipLimit=2.5, tileGridSize=(12, 12))
        enhanced = clahe.apply(gray)
        _, thresh2 = cv2.threshold(
            enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        processed_versions.append(('otsu', thresh2))

        # Method 3: Bilateral filter + adaptive threshold
        filtered = cv2.bilateralFilter(gray, 9, 75, 75)
        thresh3 = cv2.adaptiveThreshold(filtered, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                        cv2.THRESH_BINARY, 15, 10)
        processed_versions.append(('bilateral', thresh3))

        # Method 4: Morphological operations
        kernel = np.ones((2, 2), np.uint8)
        morph = cv2.morphologyEx(gray, cv2.MORPH_CLOSE, kernel)
        _, thresh4 = cv2.threshold(
            morph, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        processed_versions.append(('morphological', thresh4))

        # Choose the best version based on text characteristics
        best_version = processed_versions[0][1]  # Default to first

        # For now, return the CLAHE version as it generally works well
        # In a production system, you might want to test each version with OCR
        for name, version in processed_versions:
            if name == 'clahe_adaptive' and language_script in ['devanagari', 'bengali', 'gujarati']:
                best_version = version
                break
            elif name == 'otsu':
                best_version = version
                break

        # Final cleanup
        kernel = np.ones((1, 1), np.uint8)
        cleaned = cv2.morphologyEx(best_version, cv2.MORPH_CLOSE, kernel)

        # Convert back to PIL Image
        return Image.fromarray(cleaned)

    except Exception as e:
        st.warning(
            f"Enhanced image preprocessing failed: {str(e)}, using original image")
        return image


def extract_text_with_enhanced_indic_ocr(images_data: List[Tuple[int, Image.Image]],
                                         selected_languages: List[str],
                                         preprocessing=True) -> Dict[int, Tuple[str, float]]:
    """Enhanced OCR extraction with multiple attempts for maximum accuracy"""
    results = {}

    try:
        # Create language string for Tesseract
        lang_codes = [INDIC_LANGUAGES.get(lang, 'eng')
                      for lang in selected_languages]
        lang_string = '+'.join(lang_codes)

        # Determine script type from languages
        script_type = 'devanagari'  # default
        if any(lang in ['bengali'] for lang in selected_languages):
            script_type = 'bengali'
        elif any(lang in ['tamil', 'malayalam', 'kannada', 'telugu'] for lang in selected_languages):
            script_type = 'south_indian'

        # Process images with enhanced accuracy focus
        for page_num, image in images_data:
            try:
                best_text = ""
                best_confidence = 0

                # Try multiple preprocessing approaches
                preprocessing_methods = []

                if preprocessing:
                    # Enhanced preprocessing
                    processed_image = enhanced_preprocess_image_for_indic_ocr(
                        image, script_type)
                    preprocessing_methods.append(('enhanced', processed_image))

                    # Original image as backup
                    preprocessing_methods.append(('original', image))
                else:
                    preprocessing_methods.append(('original', image))

                # Try different combinations of settings
                for preprocess_name, proc_image in preprocessing_methods:
                    # Try multiple PSM modes with enhanced settings
                    psm_configs = [
                        (6, '--oem 3 --psm 6 -c preserve_interword_spaces=1 -c textord_heavy_nr=1'),
                        (4, '--oem 3 --psm 4 -c preserve_interword_spaces=1'),
                        (3, '--oem 3 --psm 3 -c preserve_interword_spaces=1'),
                        (8, '--oem 3 --psm 8 -c preserve_interword_spaces=1'),
                        (13, '--oem 3 --psm 13 -c preserve_interword_spaces=1')
                    ]

                    for psm, config in psm_configs:
                        try:
                            # Get OCR result with confidence scores
                            data = pytesseract.image_to_data(proc_image, lang=lang_string,
                                                             config=config, output_type=pytesseract.Output.DICT)

                            # Calculate average confidence
                            confidences = [
                                int(conf) for conf in data['conf'] if int(conf) > 0]
                            avg_confidence = sum(
                                confidences) / len(confidences) if confidences else 0

                            # Get text
                            text = pytesseract.image_to_string(
                                proc_image, lang=lang_string, config=config)

                            # Enhanced scoring that considers both confidence and text quality
                            quality_score = avg_confidence
                            if text.strip():
                                # Bonus for longer text (usually indicates better extraction)
                                length_bonus = min(10, len(text.strip()) / 50)
                                quality_score += length_bonus

                                # Bonus for proper word structure
                                words = text.split()
                                if len(words) > 2:
                                    quality_score += 5

                                # Penalty for too many special characters (usually indicates poor OCR)
                                special_char_ratio = sum(
                                    1 for c in text if not c.isalnum() and not c.isspace()) / len(text)
                                if special_char_ratio > 0.3:
                                    quality_score -= 10

                            # Keep the result with highest quality score
                            if quality_score > best_confidence and text.strip():
                                best_confidence = quality_score
                                best_text = text
                                st.write(
                                    f"📄 Page {page_num + 1}: Found better result with {preprocess_name} preprocessing, PSM {psm} (score: {quality_score:.1f})")

                        except Exception as psm_error:
                            continue

                # Cap confidence at reasonable maximum
                final_confidence = min(95.0, best_confidence)
                results[page_num] = (best_text, final_confidence)

                # Clean up processed images
                for _, proc_img in preprocessing_methods:
                    if proc_img != image:  # Don't delete original
                        del proc_img

                if get_memory_usage() > memory_threshold:
                    cleanup_memory()

            except Exception as e:
                st.warning(
                    f"Enhanced OCR extraction failed for page {page_num + 1}: {str(e)}")
                results[page_num] = ("", 0)

    except Exception as e:
        st.error(f"Enhanced batch OCR extraction failed: {str(e)}")

    return results


def process_page_batch_accuracy_focused(pdf_document, page_range: range, extraction_method: str,
                                        selected_languages: List[str], enable_preprocessing: bool,
                                        gemini_extraction_type: str = "general") -> Dict[int, Dict]:
    """Process a batch of pages with accuracy as the primary focus"""
    batch_results = {}

    try:
        st.info(
            f"🎯 Processing pages {page_range.start + 1}-{page_range.stop} with {extraction_method} method (accuracy-focused)")

        # For Gemini-only mode, skip regular text extraction entirely
        if extraction_method == "gemini":
            images_for_gemini = []

            # Convert all pages to high-quality images for Gemini processing
            for page_num in page_range:
                try:
                    page = pdf_document[page_num]

                    # High-resolution image extraction for maximum accuracy
                    # Higher resolution for better accuracy
                    mat = fitz.Matrix(3.0, 3.0)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))

                    images_for_gemini.append((page_num, img))

                    batch_results[page_num] = {
                        'text': '',
                        'method': 'pending_gemini',
                        'confidence': 0
                    }

                except Exception as e:
                    batch_results[page_num] = {
                        'text': f"Error preparing page for Gemini: {str(e)}",
                        'method': 'error',
                        'confidence': 0
                    }

            # Process all pages with Gemini
            if images_for_gemini:
                st.info(
                    f"🤖 Processing {len(images_for_gemini)} pages with Gemini AI...")
                gemini_results = extract_text_with_gemini_accuracy_focused(
                    images_for_gemini, selected_languages, gemini_extraction_type)

                # Update results
                for page_num, (text, confidence) in gemini_results.items():
                    batch_results[page_num] = {
                        'text': text,
                        'method': 'gemini',
                        'confidence': confidence
                    }

            # Clean up
            for _, img in images_for_gemini:
                del img
            del images_for_gemini

        else:
            # For other methods, use the original logic but with enhancements
            images_for_ai = []

            # First pass: extract regular text and prepare images for AI processing
            for page_num in page_range:
                try:
                    page = pdf_document[page_num]

                    # Try regular text extraction first (except for gemini-only mode)
                    page_text = page.get_text()

                    # Enhanced criteria for when to use AI extraction
                    needs_ai_extraction = (
                        not page_text.strip() or
                        len(page_text.strip()) < 20 or  # Stricter threshold
                        extraction_method in ["tesseract_ocr", "hybrid"]
                    )

                    if needs_ai_extraction:
                        # Convert page to high-quality image
                        mat = fitz.Matrix(3.0, 3.0)  # Higher resolution
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))

                        images_for_ai.append((page_num, img))

                        batch_results[page_num] = {
                            'text': page_text if page_text.strip() else '',
                            'method': 'pending_ai',
                            'confidence': 50.0 if page_text.strip() else 0
                        }
                    else:
                        batch_results[page_num] = {
                            'text': page_text,
                            'method': 'regular',
                            'confidence': 95.0
                        }

                except Exception as e:
                    batch_results[page_num] = {
                        'text': f"Error processing page: {str(e)}",
                        'method': 'error',
                        'confidence': 0
                    }

            # AI extraction for pages that need it
            if images_for_ai:
                if extraction_method in ["gemini", "hybrid"]:
                    st.info(
                        f"🤖 Processing {len(images_for_ai)} pages with Gemini AI...")
                    gemini_results = extract_text_with_gemini_accuracy_focused(
                        images_for_ai, selected_languages, gemini_extraction_type)

                    # Update results
                    for page_num, (text, confidence) in gemini_results.items():
                        if text.strip():
                            batch_results[page_num] = {
                                'text': text,
                                'method': 'gemini',
                                'confidence': confidence
                            }

                if extraction_method in ["tesseract_ocr", "hybrid"]:
                    st.info(
                        f"🔍 Processing {len(images_for_ai)} pages with Enhanced OCR...")
                    ocr_results = extract_text_with_enhanced_indic_ocr(
                        images_for_ai, selected_languages, enable_preprocessing)

                    # Update results for OCR
                    for page_num, (text, confidence) in ocr_results.items():
                        if extraction_method == "tesseract_ocr":
                            # Pure OCR mode
                            batch_results[page_num] = {
                                'text': text,
                                'method': 'enhanced_ocr',
                                'confidence': confidence
                            }
                        elif extraction_method == "hybrid":
                            # Hybrid mode: combine or choose best result
                            existing_text = batch_results[page_num]['text']
                            existing_confidence = batch_results[page_num]['confidence']

                            # Choose the result with higher confidence
                            if confidence > existing_confidence and text.strip():
                                batch_results[page_num] = {
                                    'text': text,
                                    'method': 'enhanced_ocr',
                                    'confidence': confidence
                                }
                            elif existing_text.strip() and text.strip():
                                # If both have decent confidence, combine them
                                combined_text = existing_text + "\n\n--- OCR Enhancement ---\n\n" + text
                                batch_results[page_num] = {
                                    'text': combined_text,
                                    'method': 'hybrid',
                                    'confidence': max(existing_confidence, confidence)
                                }

            # Clean up images
            for _, img in images_for_ai:
                del img
            del images_for_ai

        # Memory cleanup after batch processing
        if get_memory_usage() > memory_threshold:
            cleanup_memory()

    except Exception as e:
        st.error(f"Batch processing failed: {str(e)}")

    return batch_results


def merge_extraction_results(results_dict: Dict[int, Dict]) -> str:
    """Merge extracted text from all pages with enhanced formatting"""
    try:
        merged_text = []

        # Sort pages by number
        sorted_pages = sorted(results_dict.keys())

        for page_num in sorted_pages:
            page_data = results_dict[page_num]
            text = page_data.get('text', '').strip()
            method = page_data.get('method', 'unknown')
            confidence = page_data.get('confidence', 0)

            if text:
                # Add page header with metadata
                page_header = f"\n{'='*60}\nPAGE {page_num + 1} | Method: {method.upper()} | Confidence: {confidence:.1f}%\n{'='*60}\n"
                merged_text.append(page_header)
                merged_text.append(text)
                merged_text.append("\n")
            else:
                # Even for empty pages, add a note
                page_header = f"\n{'='*60}\nPAGE {page_num + 1} | No text extracted\n{'='*60}\n"
                merged_text.append(page_header)

        return "\n".join(merged_text)

    except Exception as e:
        st.error(f"Error merging results: {str(e)}")
        return "Error merging extraction results"


def create_enhanced_docx_with_metadata(text_content: str, extraction_metadata: Dict) -> io.BytesIO:
    """Create a professional DOCX document with extraction metadata"""
    try:
        doc = Document()

        # Title
        title = doc.add_heading('PDF Text Extraction Report', 0)
        title.alignment = 1  # Center alignment

        # Metadata section
        doc.add_heading('Extraction Details', level=1)

        # Add metadata table
        metadata_table = doc.add_table(rows=1, cols=2)
        metadata_table.style = 'Table Grid'
        hdr_cells = metadata_table.rows[0].cells
        hdr_cells[0].text = 'Property'
        hdr_cells[1].text = 'Value'

        # Add metadata rows
        for key, value in extraction_metadata.items():
            row_cells = metadata_table.add_row().cells
            row_cells[0].text = str(key).replace('_', ' ').title()
            row_cells[1].text = str(value)

        # Add page break
        doc.add_page_break()

        # Content section
        doc.add_heading('Extracted Content', level=1)

        # Add the main content
        content_paragraph = doc.add_paragraph()
        content_paragraph.add_run(text_content)

        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        return doc_buffer

    except Exception as e:
        st.error(f"Error creating DOCX: {str(e)}")
        return None


def main():
    """Main Streamlit application with enhanced UI and persistent state"""

    # Initialize session state for persistence
    if 'extraction_complete' not in st.session_state:
        st.session_state.extraction_complete = False
    if 'extracted_results' not in st.session_state:
        st.session_state.extracted_results = {}
    if 'merged_text' not in st.session_state:
        st.session_state.merged_text = ""
    if 'extraction_metadata' not in st.session_state:
        st.session_state.extraction_metadata = {}
    if 'processing_stats' not in st.session_state:
        st.session_state.processing_stats = {}

    # Custom CSS for better appearance
    st.markdown("""
    <style>
    .main-header {
        text-align: center;
        color: #1f77b4;
        margin-bottom: 30px;
    }
    .stProgress > div > div > div > div {
        background-color: #1f77b4;
    }
    .status-box {
        padding: 10px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
    }
    .error-box {
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        color: #721c24;
    }
    .info-box {
        background-color: #d1ecf1;
        border: 1px solid #bee5eb;
        color: #0c5460;
    }
    </style>
    """, unsafe_allow_html=True)

    # Main header
    st.markdown('<h1 class="main-header">🚀 Enhanced Large PDF Text Extractor with AI</h1>',
                unsafe_allow_html=True)

    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")

        # Show current status
        if st.session_state.get('extraction_complete', False):
            st.success("✅ Extraction Complete")
            if st.button("🗑️ Clear Results", help="Clear current results to process a new file"):
                clear_extraction_state()
                st.rerun()

        # API Key section
        st.subheader("🔑 AI Configuration")
        gemini_api_key = st.text_input("Gemini API Key", type="password",
                                       help="Enter your Google Gemini API key for AI-powered extraction")

        # Language selection
        st.subheader("🌐 Language Selection")
        available_languages = list(INDIC_LANGUAGES.keys())
        selected_languages = st.multiselect(
            "Select languages for extraction",
            available_languages,
            default=['english', 'hindi'],
            help="Choose the languages present in your PDF"
        )

        # Extraction method
        st.subheader("🎯 Extraction Method")
        extraction_methods = {
            "hybrid": "🔀 Hybrid (Regular + AI)",
            "gemini": "🤖 Gemini AI Only",
            "tesseract_ocr": "🔍 Enhanced OCR Only",
            "regular": "📄 Regular PDF Text Only"
        }

        extraction_method = st.selectbox(
            "Choose extraction method",
            options=list(extraction_methods.keys()),
            format_func=lambda x: extraction_methods[x],
            index=0,
            help="Hybrid is recommended for best results"
        )

        # Gemini-specific options
        if extraction_method in ["hybrid", "gemini"]:
            st.subheader("🤖 Gemini AI Options")
            gemini_extraction_types = {
                "general": "📋 General Document",
                "academic": "🎓 Academic/Research Paper",
                "indic_specialized": "🕉️ Indic Languages Specialized",
                "handwritten": "✍️ Handwritten Text"
            }

            gemini_extraction_type = st.selectbox(
                "Document Type",
                options=list(gemini_extraction_types.keys()),
                format_func=lambda x: gemini_extraction_types[x],
                help="Choose the type that best matches your document"
            )
        else:
            gemini_extraction_type = "general"

        # Processing options
        st.subheader("⚙️ Processing Options")
        enable_preprocessing = st.checkbox("Enable Image Preprocessing", value=True,
                                           help="Applies image enhancement for better OCR results")

        batch_size = st.slider("Batch Size", min_value=1, max_value=20, value=5,
                               help="Number of pages to process at once")

        # Advanced options
        with st.expander("🔧 Advanced Options"):
            max_pages = st.number_input("Max Pages to Process", min_value=1, max_value=1000,
                                        value=100, help="Limit processing to first N pages")

            enable_debug = st.checkbox("Enable Debug Mode", value=False,
                                       help="Show detailed processing information")

    # Main content area
    st.header("📁 File Upload")

    uploaded_file = st.file_uploader(
        "Choose a PDF file",
        type="pdf",
        help="Upload the PDF file you want to extract text from"
    )

    if uploaded_file is not None:
        # Validate API key for AI methods
        if extraction_method in ["hybrid", "gemini"] and not gemini_api_key:
            st.error(
                "🔑 Gemini API key is required for AI-powered extraction methods!")
            st.stop()

        # Setup Gemini if needed
        if extraction_method in ["hybrid", "gemini"]:
            if not setup_gemini(gemini_api_key):
                st.error(
                    "❌ Failed to setup Gemini API. Please check your API key.")
                st.stop()
            else:
                st.success("✅ Gemini API setup successful!")

        # Load PDF
        try:
            pdf_bytes = uploaded_file.read()
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            total_pages = len(pdf_document)

            st.success(
                f"📖 PDF loaded successfully! Total pages: {total_pages}")

            # Limit pages if specified
            pages_to_process = min(total_pages, max_pages)
            if pages_to_process < total_pages:
                st.warning(
                    f"⚠️ Processing limited to first {pages_to_process} pages")

        except Exception as e:
            st.error(f"❌ Error loading PDF: {str(e)}")
            st.stop()

        # Display file information
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("File Size", f"{len(pdf_bytes) / (1024*1024):.1f} MB")
        with col2:
            st.metric("Total Pages", total_pages)
        with col3:
            st.metric("Processing Method",
                      extraction_methods[extraction_method])

        # Processing section
        st.header("🚀 Processing")

        # Create columns for better layout
        col1, col2 = st.columns([3, 1])

        with col1:
            start_extraction = st.button("Start Extraction", type="primary",
                                         disabled=st.session_state.extraction_complete)

        with col2:
            if st.session_state.extraction_complete:
                if st.button("🔄 Process New File", type="secondary"):
                    # Clear session state for new processing
                    clear_extraction_state()
                    st.rerun()

        # Show processing status
        if st.session_state.extraction_complete:
            st.success("✅ Extraction completed! Downloads are ready below.")

        # Process only if button clicked and not already complete
        if start_extraction and not st.session_state.extraction_complete:
            start_time = time.time()

            # Initialize progress tracking
            progress_bar = st.progress(0)
            status_text = st.empty()
            results_container = st.container()

            # Initialize results storage
            all_results = {}

            try:
                # Process in batches
                total_batches = (pages_to_process +
                                 batch_size - 1) // batch_size

                for batch_idx in range(total_batches):
                    start_page = batch_idx * batch_size
                    end_page = min((batch_idx + 1) *
                                   batch_size, pages_to_process)
                    page_range = range(start_page, end_page)

                    # Update progress
                    progress = (batch_idx + 1) / total_batches
                    progress_bar.progress(progress)
                    status_text.text(
                        f"Processing batch {batch_idx + 1}/{total_batches} (pages {start_page + 1}-{end_page})")

                    # Process batch
                    batch_results = process_page_batch_accuracy_focused(
                        pdf_document, page_range, extraction_method,
                        selected_languages, enable_preprocessing, gemini_extraction_type
                    )

                    # Store results
                    all_results.update(batch_results)

                    # Show intermediate results
                    if enable_debug:
                        with results_container:
                            st.write(f"**Batch {batch_idx + 1} Results:**")
                            for page_num, result in batch_results.items():
                                confidence = result.get('confidence', 0)
                                method = result.get('method', 'unknown')
                                text_length = len(result.get('text', ''))
                                st.write(
                                    f"Page {page_num + 1}: {text_length} chars, {confidence:.1f}% confidence ({method})")

                # Store results in session state instead of local variables
                st.session_state.extracted_results = all_results
                st.session_state.merged_text = merge_extraction_results(
                    all_results)

                # Create and store extraction metadata
                processing_time = time.time() - start_time
                st.session_state.extraction_metadata = {
                    'extraction_method': extraction_methods[extraction_method],
                    'languages_selected': ', '.join(selected_languages),
                    'total_pages_processed': pages_to_process,
                    'processing_time_seconds': f"{processing_time:.2f}",
                    'total_characters_extracted': len(st.session_state.merged_text),
                    'gemini_extraction_type': gemini_extraction_types.get(gemini_extraction_type, 'N/A') if extraction_method in ["hybrid", "gemini"] else 'N/A',
                    'preprocessing_enabled': enable_preprocessing,
                    'extraction_timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }

                # Store processing stats
                avg_confidence = sum(r.get('confidence', 0) for r in all_results.values(
                )) / len(all_results) if all_results else 0
                method_counts = {}
                for result in all_results.values():
                    method = result.get('method', 'unknown')
                    method_counts[method] = method_counts.get(method, 0) + 1

                st.session_state.processing_stats = {
                    'pages_processed': pages_to_process,
                    'processing_time': processing_time,
                    'total_characters': len(st.session_state.merged_text),
                    'avg_confidence': avg_confidence,
                    'method_counts': method_counts
                }

                # Mark extraction as complete
                st.session_state.extraction_complete = True

                # Complete processing
                progress_bar.progress(1.0)
                status_text.text("Processing complete! Results saved.")

                # Force a rerun to show results
                st.rerun()

            except Exception as e:
                st.error(f"❌ Processing failed: {str(e)}")
                if enable_debug:
                    st.exception(e)
            finally:
                # Cleanup
                if 'pdf_document' in locals():
                    pdf_document.close()
                cleanup_memory()

        # Display results section (always show if extraction is complete)
        if st.session_state.extraction_complete and st.session_state.merged_text:
            st.header("📊 Extraction Results")

            # Display stats from session state
            stats = st.session_state.processing_stats
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Pages Processed", stats.get('pages_processed', 0))
            with col2:
                st.metric("Processing Time",
                          f"{stats.get('processing_time', 0):.1f}s")
            with col3:
                st.metric("Characters Extracted",
                          f"{stats.get('total_characters', 0):,}")
            with col4:
                st.metric("Avg Confidence",
                          f"{stats.get('avg_confidence', 0):.1f}%")

            # Results breakdown by method
            method_counts = stats.get('method_counts', {})
            if len(method_counts) > 1:
                st.subheader("📈 Processing Method Breakdown")
                for method, count in method_counts.items():
                    st.write(
                        f"**{method.replace('_', ' ').title()}**: {count} pages")

            # Text preview
            st.subheader("👁️ Text Preview")
            preview_text = st.session_state.merged_text[:2000] + \
                ("..." if len(st.session_state.merged_text) > 2000 else "")
            st.text_area("Extracted Text Preview", preview_text, height=300)

            # Persistent download section
            st.header("💾 Download Options")

            # Add helpful message
            st.info(
                "💡 **Download Tip**: You can download multiple formats! Each download won't reset the app.")

            col1, col2, col3 = st.columns(3)

            with col1:
                # Plain text download with unique key
                st.download_button(
                    label="📄 Download as TXT",
                    data=st.session_state.merged_text.encode('utf-8'),
                    file_name=f"extracted_text_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    key="download_txt"  # Unique key prevents conflicts
                )

            with col2:
                # DOCX download with unique key
                try:
                    docx_buffer = create_enhanced_docx_with_metadata(
                        st.session_state.merged_text,
                        st.session_state.extraction_metadata
                    )
                    if docx_buffer:
                        st.download_button(
                            label="📝 Download as DOCX",
                            data=docx_buffer.getvalue(),
                            file_name=f"extracted_text_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_docx"  # Unique key
                        )
                    else:
                        st.error("Error creating DOCX file")
                except Exception as e:
                    st.error(f"Error creating DOCX: {str(e)}")

            with col3:
                # JSON download for debug (if enabled)
                if enable_debug:
                    json_data = {
                        'metadata': st.session_state.extraction_metadata,
                        'results': {str(k): v for k, v in st.session_state.extracted_results.items()},
                        'merged_text': st.session_state.merged_text
                    }

                    st.download_button(
                        label="🔧 Download Debug Data (JSON)",
                        data=json.dumps(json_data, indent=2,
                                        ensure_ascii=False).encode('utf-8'),
                        file_name=f"extraction_debug_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json",
                        key="download_json"  # Unique key
                    )

            # Additional download status
            st.success(
                "🎉 **All downloads are ready!** Click any download button above. The app won't reset between downloads.")

    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #666;'>
        <p>🚀 Enhanced PDF Text Extractor | Supports 12+ Languages | AI-Powered Accuracy</p>
        <p>💡 For best results with Indic languages, use Hybrid or Gemini AI methods</p>
    </div>
    """, unsafe_allow_html=True)


def clear_extraction_state():
    """Helper function to clear extraction state"""
    keys_to_clear = [
        'extraction_complete', 'extracted_results', 'merged_text',
        'extraction_metadata', 'processing_stats'
    ]
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]


if __name__ == "__main__":
    # Set up logging
    logging.basicConfig(level=logging.INFO)

    # Configure Tesseract path if needed (uncomment and modify as per your system)
    # pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'  # Linux/Mac
    # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Windows

    main()
